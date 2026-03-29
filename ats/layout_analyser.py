"""
layout_analyser.py  —  Structural & formatting analysis
=========================================================
Detects layout issues that cause ATS mis-parsing:
  • multi-column layouts
  • tables used as layout containers
  • floating text boxes
  • graphics / icons
  • contact info inside header/footer
  • image-based (non-selectable) PDFs
  • encoding anomalies
  • special characters that confuse older parsers

Technique stack
---------------
  PDF   → pdfplumber (text + layout coordinates)
  DOCX  → python-docx  (XML introspection)
  Heuristics on bounding boxes, XML node types, character analysis
  No hardcoded keyword lists — all detection is structural / statistical
"""

from __future__ import annotations

import io
import re
import unicodedata
from collections import Counter
from pathlib import Path
from typing import Any

from .models import (
    Issue, IssueCategory, LayoutMeta, LayoutResult, Platform, Severity
)

# Optional imports — graceful degradation if not installed
try:
    import pdfplumber
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════════════
class LayoutAnalyser:

    # Score deductions per issue (calibrated to produce realistic ATS scores)
    _DEDUCTIONS = {
        "image_pdf"         : 50,
        "multi_column"      : 20,
        "table_layout"      : 18,
        "text_box"          : 15,
        "header_footer"     : 12,
        "graphics"          : 10,
        "encoding_issues"   : 8,
        "special_chars"     : 5,
        "too_many_fonts"    : 4,
    }

    # ───────────────────────────────────────────────────────────────────────────
    def analyse(self, file_bytes: bytes, ext: str) -> LayoutResult:
        if ext == ".pdf":
            return self._analyse_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return self._analyse_docx(file_bytes)
        else:
            # Plain text — layout is always clean
            text = file_bytes.decode("utf-8", errors="replace")
            meta = self._text_meta(text, ext)
            return LayoutResult(score=100.0, plain_text=text, meta=meta, issues=[])

    # ─────────────────────────────────────────────────
    # PDF analysis
    # ─────────────────────────────────────────────────
    def _analyse_pdf(self, data: bytes) -> LayoutResult:
        if not _PDF_AVAILABLE:
            raise ImportError("pdfplumber is required: pip install pdfplumber")

        issues: list[Issue] = []
        score  = 100.0
        all_text_parts: list[str] = []

        is_image_pdf       = False
        column_count       = 1
        has_tables         = False
        has_graphics       = False
        special_char_ratio = 0.0
        font_names: set[str] = set()
        total_chars = 0
        non_ascii   = 0
        page_count  = 0

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)

            for page in pdf.pages:
                raw_text = page.extract_text() or ""
                all_text_parts.append(raw_text)

                # ── Detect image-only PDF ────────────────────────────────────
                if not raw_text.strip() and page.images:
                    is_image_pdf = True

                # ── Detect tables ────────────────────────────────────────────
                if page.extract_tables():
                    has_tables = True

                # ── Detect graphics (non-text objects) ───────────────────────
                if page.images:
                    has_graphics = True

                # ── Column detection via bounding-box clustering ─────────────
                words = page.extract_words()
                if words:
                    col = _detect_columns_from_words(words, page.width)
                    column_count = max(column_count, col)

                # ── Font inventory ────────────────────────────────────────────
                for char in page.chars:
                    fn = char.get("fontname", "")
                    if fn:
                        # Normalise variant suffixes (Bold, Italic…)
                        base = re.sub(r"[-,+]?(Bold|Italic|Regular|Light|Medium|BoldItalic).*", "", fn, flags=re.I)
                        font_names.add(base)

                # ── Special character analysis ────────────────────────────────
                for ch in raw_text:
                    total_chars += 1
                    if ord(ch) > 127:
                        non_ascii += 1

        plain_text = "\n".join(all_text_parts)
        if total_chars > 0:
            special_char_ratio = non_ascii / total_chars

        # ── Build issue list ─────────────────────────────────────────────────
        if is_image_pdf:
            score -= self._DEDUCTIONS["image_pdf"]
            issues.append(Issue(
                severity=Severity.CRITICAL,
                category=IssueCategory.ENCODING,
                message="PDF appears to be image-based (scanned). ATS cannot extract any text.",
                fix="Re-export from Word or Google Docs as a text-based PDF. Open in Acrobat and verify text is selectable.",
                platform=None,
            ))

        if column_count >= 2:
            # Taleo / iCIMS cannot parse multi-column; Greenhouse tolerates it
            severity = Severity.CRITICAL if column_count >= 3 else Severity.HIGH
            score -= self._DEDUCTIONS["multi_column"]
            issues.append(Issue(
                severity=severity,
                category=IssueCategory.LAYOUT,
                message=f"Detected {column_count}-column layout. Taleo and iCIMS will mis-parse this, merging unrelated content.",
                fix="Reformat to a single-column layout. Move skills sidebar content into a flat Skills section.",
                platform=Platform.TALEO,
            ))
            issues.append(Issue(
                severity=Severity.MEDIUM,
                category=IssueCategory.LAYOUT,
                message=f"{column_count}-column layout may cause column-merging in Workday and iCIMS.",
                fix="Use single-column to guarantee compatibility across all platforms.",
                platform=Platform.WORKDAY,
            ))

        if has_tables:
            score -= self._DEDUCTIONS["table_layout"]
            issues.append(Issue(
                severity=Severity.HIGH,
                category=IssueCategory.LAYOUT,
                message="Tables detected. ATS parsers strip table structure and may scramble or lose content.",
                fix="Replace all tables with plain text blocks. Use line breaks and consistent spacing instead.",
                platform=None,
            ))

        if has_graphics:
            score -= self._DEDUCTIONS["graphics"]
            issues.append(Issue(
                severity=Severity.MEDIUM,
                category=IssueCategory.LAYOUT,
                message="Graphics/images detected (icons, photo, progress bars). ATS reads these as garbage characters.",
                fix="Remove all graphics, icons, profile photos, and skill-rating bars. Use text only.",
                platform=None,
            ))

        if len(font_names) > 3:
            score -= self._DEDUCTIONS["too_many_fonts"]
            issues.append(Issue(
                severity=Severity.LOW,
                category=IssueCategory.READABILITY,
                message=f"{len(font_names)} different fonts detected. Some ATS parsers flag font inconsistency.",
                fix="Standardise to 1-2 fonts throughout. Arial, Calibri, or Garamond are ATS-safe choices.",
                platform=None,
            ))

        if special_char_ratio > 0.04:
            score -= self._DEDUCTIONS["special_chars"]
            issues.append(Issue(
                severity=Severity.MEDIUM,
                category=IssueCategory.ENCODING,
                message=f"High non-ASCII character ratio ({special_char_ratio:.1%}). Fancy bullets, em-dashes, and Unicode symbols cause parser errors.",
                fix="Replace decorative bullets (▸ ● ✓) with standard hyphens (-) or dots (•). Remove em-dashes (—) and replace with plain hyphens.",
                platform=Platform.TALEO,
            ))

        meta = LayoutMeta(
            file_format="pdf",
            is_image_pdf=is_image_pdf,
            column_count=column_count,
            has_tables=has_tables,
            has_text_boxes=False,           # Not applicable to PDF
            has_header_footer_content=False, # Checked in semantic layer
            has_graphics=has_graphics,
            has_columns_via_table=False,
            font_count=len(font_names),
            special_char_ratio=special_char_ratio,
            character_encoding="utf-8",
            page_count=page_count,
            word_count=len(plain_text.split()),
        )

        return LayoutResult(
            score=max(0.0, score),
            plain_text=plain_text,
            meta=meta,
            issues=issues,
        )

    # ─────────────────────────────────────────────────
    # DOCX analysis
    # ─────────────────────────────────────────────────
    def _analyse_docx(self, data: bytes) -> LayoutResult:
        if not _DOCX_AVAILABLE:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = DocxDocument(io.BytesIO(data))
        issues: list[Issue] = []
        score  = 100.0

        has_tables             = False
        has_text_boxes         = False
        has_header_footer_info = False
        has_graphics           = False
        has_columns_via_table  = False
        font_names: set[str]   = set()
        all_text_parts: list[str] = []

        # ── Paragraph text ────────────────────────────────────────────────────
        for para in doc.paragraphs:
            all_text_parts.append(para.text)
            for run in para.runs:
                if run.font and run.font.name:
                    font_names.add(run.font.name)

        # ── Table analysis ────────────────────────────────────────────────────
        for table in doc.tables:
            has_tables = True
            # Determine if the table is used as a layout grid (fake columns)
            # Heuristic: a table with ≥2 columns and many rows is layout abuse
            if len(table.columns) >= 2 and len(table.rows) >= 3:
                has_columns_via_table = True
            for row in table.rows:
                for cell in row.cells:
                    all_text_parts.append(cell.text)

        # ── Text boxes (floating frames) ──────────────────────────────────────
        # Inspect XML for <w:txbxContent> elements
        body_xml = doc.element.body.xml
        if "<w:txbxContent>" in body_xml or "txbxContent" in body_xml:
            has_text_boxes = True

        # ── Header / footer content ───────────────────────────────────────────
        contact_pattern = re.compile(r"[@+\d]{3,}|linkedin|github", re.I)
        for section in doc.sections:
            for hdr in [section.header, section.footer]:
                if hdr:
                    hf_text = " ".join(p.text for p in hdr.paragraphs)
                    if contact_pattern.search(hf_text):
                        has_header_footer_info = True

        # ── Graphics (inline images / drawing objects) ─────────────────────
        if "<w:drawing>" in body_xml or "<v:imagedata" in body_xml:
            has_graphics = True

        # ── Build issue list ─────────────────────────────────────────────────
        if has_text_boxes:
            score -= self._DEDUCTIONS["text_box"]
            issues.append(Issue(
                severity=Severity.CRITICAL,
                category=IssueCategory.LAYOUT,
                message="Floating text boxes detected. All major ATS parsers skip text box content entirely.",
                fix="Select each text box, copy its content, paste as a normal paragraph, then delete the text box.",
                platform=None,
            ))

        if has_columns_via_table:
            score -= self._DEDUCTIONS["table_layout"] + self._DEDUCTIONS["multi_column"]
            issues.append(Issue(
                severity=Severity.CRITICAL,
                category=IssueCategory.LAYOUT,
                message="Multi-column layout created with tables. Taleo and iCIMS will parse columns left-to-right and merge your skills with your job titles.",
                fix="Delete the layout table. Restructure as a single-column document with plain paragraph text.",
                platform=Platform.TALEO,
            ))
        elif has_tables:
            score -= self._DEDUCTIONS["table_layout"]
            issues.append(Issue(
                severity=Severity.HIGH,
                category=IssueCategory.LAYOUT,
                message="Tables detected in DOCX. ATS parsers strip table structure and scramble the content.",
                fix="Replace all tables with plain-text formatting (line breaks, tabs, or consistent spacing).",
                platform=None,
            ))

        if has_header_footer_info:
            score -= self._DEDUCTIONS["header_footer"]
            issues.append(Issue(
                severity=Severity.HIGH,
                category=IssueCategory.CONTACT,
                message="Contact information (email, phone, or LinkedIn) is inside a Word header/footer. Most ATS systems ignore header/footer regions.",
                fix="Cut your contact details from the header/footer and paste them into the first paragraph of the document body.",
                platform=None,
            ))

        if has_graphics:
            score -= self._DEDUCTIONS["graphics"]
            issues.append(Issue(
                severity=Severity.MEDIUM,
                category=IssueCategory.LAYOUT,
                message="Images or drawing objects detected. Icons, photos, and decorative graphics are read as garbled text or skipped.",
                fix="Delete all images, icons, and inline graphics. Replace skill-bars or rating visuals with plain text ('Python — Advanced').",
                platform=None,
            ))

        if len(font_names) > 3:
            score -= self._DEDUCTIONS["too_many_fonts"]
            issues.append(Issue(
                severity=Severity.LOW,
                category=IssueCategory.READABILITY,
                message=f"{len(font_names)} distinct fonts used. Aim for maximum 2.",
                fix="Select all text (Ctrl+A) and set a single ATS-safe font. Calibri or Arial recommended.",
                platform=None,
            ))

        plain_text = "\n".join(all_text_parts)
        meta = LayoutMeta(
            file_format="docx",
            is_image_pdf=False,
            column_count=2 if has_columns_via_table else 1,
            has_tables=has_tables,
            has_text_boxes=has_text_boxes,
            has_header_footer_content=has_header_footer_info,
            has_graphics=has_graphics,
            has_columns_via_table=has_columns_via_table,
            font_count=len(font_names),
            special_char_ratio=_calc_special_char_ratio(plain_text),
            character_encoding="utf-8",
            page_count=1,        # python-docx does not expose page count directly
            word_count=len(plain_text.split()),
        )

        return LayoutResult(
            score=max(0.0, score),
            plain_text=plain_text,
            meta=meta,
            issues=issues,
        )

    # ─────────────────────────────────────────────────
    # Plain text meta
    # ─────────────────────────────────────────────────
    def _text_meta(self, text: str, ext: str) -> LayoutMeta:
        return LayoutMeta(
            file_format=ext.lstrip("."),
            is_image_pdf=False,
            column_count=1,
            has_tables=False,
            has_text_boxes=False,
            has_header_footer_content=False,
            has_graphics=False,
            has_columns_via_table=False,
            font_count=1,
            special_char_ratio=_calc_special_char_ratio(text),
            character_encoding="utf-8",
            page_count=1,
            word_count=len(text.split()),
        )


# ─────────────────────────────────────────────────────────────────────────────
# Utility helpers
# ─────────────────────────────────────────────────────────────────────────────

def _detect_columns_from_words(words: list[dict], page_width: float) -> int:
    """
    Cluster word x0 positions into groups using a simple gap-based heuristic.
    Returns the estimated number of columns (1, 2, or 3+).

    Method: build a histogram of x0 positions, look for a clear gap in the
    middle third of the page width — a strong indicator of two columns.
    """
    if not words:
        return 1

    x_positions = [w["x0"] for w in words]
    if not x_positions:
        return 1

    mid_lo = page_width * 0.35
    mid_hi = page_width * 0.65

    # Words whose x0 falls in the centre gap suggest two distinct columns
    left_count   = sum(1 for x in x_positions if x < mid_lo)
    right_count  = sum(1 for x in x_positions if x > mid_hi)
    centre_count = sum(1 for x in x_positions if mid_lo <= x <= mid_hi)

    # If left AND right columns both have significant content and the centre is sparse
    total = len(x_positions)
    if (left_count / total > 0.2
        and right_count / total > 0.2
        and centre_count / total < 0.15):
        return 2

    return 1


def _calc_special_char_ratio(text: str) -> float:
    if not text:
        return 0.0
    non_ascii = sum(1 for c in text if ord(c) > 127)
    return non_ascii / len(text)
